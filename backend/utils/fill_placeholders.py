import os
import re
from docx import Document

DEFAULT_TEMPLATE = os.path.join(
    os.path.dirname(__file__),
    "..",
    "sample_docs",
    "Postmoney_Safe_Template.docx"
)
DEFAULT_TEMPLATE = os.path.abspath(DEFAULT_TEMPLATE)

DOLLAR_BLANK_RE = re.compile(r'(\$\s*)\[\s*_+\s*\]')

def _replace_basic_tokens(text: str, mapping: dict, keys: list) -> str:
    """Replace {{key}}, [key], and [key with spaces] straightforwardly."""
    new_text = text
    for key in keys:
        val = str(mapping.get(key, ""))
        if not val:
            continue
        patterns = [
            "{{" + key + "}}",
            "[" + key + "]",
            "[" + key.replace("_", " ") + "]",
        ]
        for pattern in patterns:
            if pattern in new_text:
                new_text = new_text.replace(pattern, val)
    return new_text

def _replace_post_money_cap(text: str, value: str) -> str:
    # Replace the $[____] following "Post-Money Valuation Cap"
    pattern = re.compile(
        r'(Post[\-\s]Money\s+Valuation\s+Cap[^$]*\$\s*)\[\s*_+\s*\]',
        flags=re.IGNORECASE | re.UNICODE
    )
    return pattern.sub(rf'\1{value}', text)

def _replace_purchase_amount(text: str, value: str) -> str:
    # Replace the $[____] in the sentence that includes (the “Purchase Amount”) ... of $[____]
    pattern = re.compile(
        r'(\(the\s*[“"]Purchase\s+Amount[”"]\).*?\bof\b\s*\$\s*)\[\s*_+\s*\]',
        flags=re.IGNORECASE | re.UNICODE | re.DOTALL
    )
    return pattern.sub(rf'\1{value}', text)

def _replace_first_dollar_blank(text: str, value: str) -> str:
    # Replace first generic $[____] with a value
    return DOLLAR_BLANK_RE.sub(rf'\1{value}', text, count=1)

def _process_block_text(text: str, mapping: dict, ordered_keys: list) -> str:
    # 1) Basic tokens
    new_text = _replace_basic_tokens(text, mapping, ordered_keys)

    # 2) Context-aware replacements if those keys exist
    if mapping.get("Post_Money_Valuation_Cap"):
        new_text = _replace_post_money_cap(new_text, str(mapping["Post_Money_Valuation_Cap"]))
    if mapping.get("Purchase_Amount"):
        new_text = _replace_purchase_amount(new_text, str(mapping["Purchase_Amount"]))

    # 3) Fallback: replace remaining generic $[____] using any unmapped Dollar_Blank_* keys
    for k in ordered_keys:
        if not k.lower().startswith("dollar_blank"):
            continue
        val = str(mapping.get(k, "")).strip()
        if not val:
            continue
        # Replace only if there is a remaining dollar blank
        if DOLLAR_BLANK_RE.search(new_text):
            new_text = _replace_first_dollar_blank(new_text, val)

    return new_text

def fill_placeholders_in_docx(input_path, placeholders, values, output_path):
    # `placeholders` is the ordered list you returned from extract_placeholders
    # `values` is a dict keyed by those placeholder names
    doc = Document(input_path)

    # Paragraphs
    for p in doc.paragraphs:
        if not p.text:
            continue
        new_text = _process_block_text(p.text, values, placeholders)
        if new_text != p.text:
            p.text = new_text  # note: resets mixed formatting in paragraph (OK for MVP)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if not cell.text:
                    continue
                cell_text = "\n".join([cp.text for cp in cell.paragraphs]) or ""
                new_text = _process_block_text(cell_text, values, placeholders)
                if new_text != cell_text:
                    # Overwrite simply (MVP)
                    for i, cp in enumerate(cell.paragraphs):
                        if i == 0:
                            cp.text = new_text
                        else:
                            cp.text = ""

    doc.save(output_path)
