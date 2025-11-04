import re
from io import BytesIO
from docx import Document

LABEL_BEFORE_DOLLAR = re.compile(
    r'(?P<label>[A-Za-z][A-Za-z0-9\s\-\u2013\u2014“”"()]+?)\s*(?:is|:)?\s*\$\s*\[\s*_+\s*\]',
    flags=re.UNICODE
)
PURCHASE_AMOUNT_OF_DOLLAR = re.compile(
    r'\(the\s*[“"](?P<label>Purchase Amount)[”"]\)\s*.*?\bof\b\s*\$\s*\[\s*_+\s*\]',
    flags=re.UNICODE | re.IGNORECASE | re.DOTALL
)
GENERIC_DOLLAR_BLANK = re.compile(r'\$\s*\[\s*_+\s*\]')

def _doc_text_with_tables(document: Document) -> str:
    parts = []
    # paragraphs
    for p in document.paragraphs:
        parts.append(p.text)
    # tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # avoid duplicate adds: join cell paragraphs
                cell_text = "\n".join([cp.text for cp in cell.paragraphs if cp.text])
                if cell_text:
                    parts.append(cell_text)
    return "\n".join([t for t in parts if t])

def _normalize_key(s: str) -> str:
    # Strip quotes/curly quotes, parentheses, punctuation around edges, then to underscores
    s = s.strip().strip('“”"()[]:. ')
    s = re.sub(r'\s+', ' ', s)
    s = s.replace('-', ' ')
    s = s.replace('\u2013', ' ').replace('\u2014', ' ')
    key = s.strip().replace(' ', '_')
    return key

def extract_placeholders(file_bytes: bytes):
    """Extract placeholders from a Word doc (supports {{ }} and [ ] formats)."""
    document = Document(BytesIO(file_bytes))
    text = []

    # Collect all text from paragraphs and tables
    for para in document.paragraphs:
        text.append(para.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text.append(cell.text)

    combined = "\n".join(text)

    # Match {{Placeholder}} and [Placeholder]
    found = re.findall(r"\{\{(.*?)\}\}|\[(.*?)\]", combined)

    # Flatten and clean both types
    placeholders = []
    for f in found:
        # f is a tuple like ('Investor Name', '') or ('', 'Company Name')
        ph = f[0].strip() if f[0].strip() else f[1].strip()
        if ph:
            placeholders.append(ph)

    # Normalize and deduplicate
    normalized = []
    seen = set()
    for ph in placeholders:
        clean = ph.strip().replace(" ", "_")
        if clean and clean not in seen:
            normalized.append(clean)
            seen.add(clean)

    # ✅ Remove any empty strings once more, just to be safe
    normalized = [ph for ph in normalized if ph]

    return normalized
